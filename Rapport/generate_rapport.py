from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

# ── Helper functions ──
def add_comment_placeholder(text):
    """Add a highlighted placeholder where a screenshot/diagram should go."""
    p = doc.add_paragraph()
    run = p.add_run(f'[INDSÆT: {text}]')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bold_italic_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    return p

def add_code_block(code_text, description=''):
    """Add a code block with monospace font and grey background indication."""
    if description:
        add_body(description)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

# ════════════════════════════════════════════════════════
#  FORSIDE (TITLE PAGE)
# ════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = t.add_run('L3ei')
run.font.size = Pt(11)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = d.add_run(datetime.date.today().strftime('%d/%m/%Y'))
run.font.size = Pt(11)

doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BetterCantine')
run.bold = True
run.font.size = Pt(28)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Kantinebestillingssystem')
run.italic = True
run.font.size = Pt(16)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Synopsis til eksamensprojekt i Programmering B')
run.italic = True
run.font.size = Pt(13)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Lavet af Aleksandar Neshev & William Holst')
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()
link_p = doc.add_paragraph()
link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = link_p.add_run('Link til program: ')
run.font.size = Pt(11)
run = link_p.add_run('[Indsæt GitHub-link / hosted link her]')
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Abstract',
    '2. Indledning & problemformulering',
    '   2.1 Problemstilling',
    '   2.2 Mål og vision',
    '   2.3 Ansvarsfordeling',
    '3. Kravspecifikation & funktionsbeskrivelse',
    '   3.1 Funktionelle krav',
    '   3.2 Ikke-funktionelle krav',
    '   3.3 Use-case diagram',
    '   3.4 Skærmbilleder',
    '4. Dokumentation af program & udvikling',
    '   4.1 Trelagsarkitektur',
    '   4.2 Database design & ER-diagram',
    '   4.3 Analyse af tredje normalform',
    '   4.4 Udvalgte use-cases med kodeeksempler',
    '      Use Case 1: Opret ny konto',
    '      Use Case 2: Log ind',
    '      Use Case 3: Se menu og bestil',
    '      Use Case 4: Administrer menu (admin)',
    '      Use Case 5: Salgsstatistik (admin)',
    '      Use Case 6: Automatisk annullering af ordrer',
    '   4.5 Trinvis udvikling',
    '5. Test af programmet',
    '   5.1 Blackbox-test',
    '   5.2 Idiottest',
    '   5.3 Fejlhåndtering (try-catch)',
    '   5.4 Brugertest',
    '6. Konklusion',
    'Bilag',
]
for item in toc_items:
    doc.add_paragraph(item)

add_body('(Opdater sidetal manuelt eller via Word-indholdsfortegnelse når rapporten er færdig.)')

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  1. ABSTRACT
# ════════════════════════════════════════════════════════
doc.add_heading('1. Abstract', level=1)

add_body(
    'Denne synopsis beskriver udviklingen af BetterCantine, et webbaseret kantinebestillingssystem '
    'bygget med Node.js, Express og Supabase (PostgreSQL). Systemet lader elever gennemse '
    'dagens menu med allergen- og halaloplysninger, bestille mad via en indkøbskurv og modtage '
    'en unik kvitteringskode til afhentning. Administratorer kan styre menuen, håndtere ordrer, '
    'sætte tidsbaserede rabatter og analysere salgsdata via et interaktivt statistikpanel med '
    'D3.js-grafer. Systemet annullerer automatisk ikke-afhentede ordrer ved kantinelukketid. '
    'Projektet demonstrerer en ren trelagsarkitektur, et relationelt databasedesign i tredje '
    'normalform, og tidsbaseret forretningslogik med reservations- og købsperioder. '
    'Synopsen gennemgår hele udviklingsprocessen fra kravspecifikation over design og '
    'implementering til test.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  2. INDLEDNING & PROBLEMFORMULERING
# ════════════════════════════════════════════════════════
doc.add_heading('2. Indledning & problemformulering', level=1)

doc.add_heading('2.1 Problemstilling', level=2)
add_body(
    'På vores skole fungerer kantinebestilling ved, at eleverne stiller sig i kø og bestiller '
    'direkte ved disken. Dette skaber lange køer, ventetid og uforudsigelighed for kantinen '
    'omkring, hvor meget mad der skal laves. Problemet er tosidet: eleverne spilder tid i '
    'pauser, og kantinepersonalet har svart ved at planlæging af madproduktionen.'
)
add_body(
    'Vi vil undersøge, om vi via programmering kan løse dette problem ved at bygge et '
    'webbaseret bestillingssystem, der giver elever mulighed for at bestille på forhånd '
    'og kantinen mulighed for at håndtere ordrer digitalt.'
)

doc.add_heading('2.2 Mål og vision', level=2)
add_body(
    'Målet med BetterCantine er at udvikle en funktionel webapplikation, der lader elever '
    'bestille mad fra skolekantinen digitalt. Systemet skal:')
add_body('- Give elever mulighed for at se dagens menu, tilføje varer til en kurv og betale')
add_body('- Generere en unik kvitteringskode, som eleven viser i kantinen for at afhente maden')
add_body('- Understøtte et reservationssystem, så elever kan reservere mad til næste dag')
add_body('- Give kantinedamen (administrator) værktøjer til at styre menu, ordrer og rabatter')
add_body(
    'Systemet bygges med Node.js og Express på serversiden, Supabase (PostgreSQL) som '
    'relationel database, og ren HTML/CSS/JavaScript på klientsiden. Vi følger en '
    'trelagsarkitektur og designer databasen i tredje normalform.'
)

doc.add_heading('2.3 Ansvarsfordeling', level=2)
add_body(
    'Projektet er lavet af Aleksandar Neshev og William Holst. '
    'Al kode er skrevet fra bunden af os selv. Vi anvender følgende tredjepartsbiblioteker, '
    'som vi importerer men ikke selv har skrevet:'
)
add_body('- Express (webserver-framework til Node.js)')
add_body('- @supabase/supabase-js (klientbibliotek til Supabase-databasen)')
add_body('- dotenv (indlæsning af miljøvariabler fra .env-fil)')
add_body('- D3.js v7 (JavaScript-bibliotek til datavisualisering, bruges til salgsstatistik-grafer)')
add_body(
    'Ingen kode er kopieret fra tidligere projekter. Databaseskemaet, al server-side logik, '
    'al klient-side JavaScript og alt HTML/CSS er skrevet af os.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  3. KRAVSPECIFIKATION & FUNKTIONSBESKRIVELSE
# ════════════════════════════════════════════════════════
doc.add_heading('3. Kravspecifikation & funktionsbeskrivelse', level=1)

doc.add_heading('3.1 Funktionelle krav', level=2)

add_bold_italic_label('Brugergodkendelse og -administration')
add_body(
    'Systemet skal give nye brugere mulighed for at registrere sig med en e-mailadresse, '
    'en adgangskode og deres fulde navn. Eksisterende brugere skal kunne logge ind. '
    'Systemet skal understøtte to brugerroller: elever (student) og administratorer (admin), '
    'med forskellige adgangsrettigheder for hver rolle.'
)

add_bold_italic_label('Gennemse menu og bestil')
add_body(
    'Eleverne skal kunne se dagens menu med alle tilgængelige food items i et '
    'kortlayout. Hvert menukort viser varenavn, beskrivelse, kategori, pris, antal tilbage '
    'og eventuel rabat. Brugeren kan tilføje og fjerne varer fra en indkøbskurv og se '
    'den samlede pris løbende.'
)

add_bold_italic_label('Reservation og direkte køb')
add_body(
    'Systemet skelner mellem to tilstande baseret på tidspunktet: '
    'en reservationsperiode (kl. 13:45 til 07:00 næste dag), hvor brugerne reserverer mad til '
    'næste dag, og en købsperiode (kl. 07:00 til 13:45), hvor brugerne køber direkte. '
    'Systemet viser automatisk den korrekte menu baseret på perioden.'
)

add_bold_italic_label('Kvitteringskode')
add_body(
    'Efter et køb eller reservation genererer systemet en unik kvitteringskode '
    '(f.eks. BC-A3F9K2), som brugeren viser i kantinen for at afhente maden. '
    'Koden genereres via en PostgreSQL-funktion i Supabase.'
)

add_bold_italic_label('Allergen- og halaloplysninger')
add_body(
    'Når en bruger holder musen over en vare, vises et detaljepanel med beskrivelse, '
    'pris, eventuel rabatpris, allergener (vist som farvede tags), halalstatus '
    '(vist som badge) og ekstra information. Disse data hentes fra food_items-tabellens '
    'udvidede kolonner (allergens, is_halal, extra_info).'
)

add_bold_italic_label('Administrative funktioner')
add_body(
    'Administratorer har adgang til et adminpanel med fire faner: '
    '"Administrer menu" (tilføj/fjern food items per dato med kategorifiltrering), '
    '"Alle ordrer" (se og marker ordrer som afhentet), '
    '"Rabatter" (angiv hvor mange styk af hver vare der sælges til rabatpris - '
    'kun tilgængeligt i kantinens åbningstid kl. 07:00-13:45), og '
    '"Statistik" (interaktivt salgsstatistikpanel).'
)

add_bold_italic_label('Salgsstatistik')
add_body(
    'Administratorer kan analysere salgsdata via et interaktivt statistikpanel bygget med D3.js. '
    'Panelet understøtter valgfri tidsperioder (1 dag, 1 uge, 1 måned, 6 måneder, 1 år) og '
    'en datovælger. Brugeren kan vælge mellem ni forskellige metrics: antal solgt, omsætning, '
    'rabat-salg, rabat-omsætning, fuld pris-salg, reservationer, direkte køb, afhentet og '
    'annulleret. Grafen viser linjediagrammer for flere dage og søjlediagrammer for enkeltdage. '
    'Hvert food item kan til-/fravælges individuelt, og en opsummering viser totaler for perioden.'
)

add_bold_italic_label('Automatisk annullering af ikke-afhentede ordrer')
add_body(
    'Når kantinen lukker kl. 13:45, annullerer systemet automatisk alle ordrer med status '
    '"paid" eller "reserved" der ikke er blevet afhentet (markeret som "picked_up"). '
    'Dette sker via en server-side timer (setInterval), der kører hvert minut og kun '
    'udløser annulleringen én gang per dag. Annullerede ordrer vises som "Annulleret" i '
    'elevens ordrehistorik.'
)

add_bold_italic_label('Tidsbaseret rabatbegrænsning')
add_body(
    'Rabatter kan kun sættes i kantinens åbningstid (kl. 07:00-13:45). Uden for dette '
    'tidsrum viser rabatsiden en besked om, at kantinen er lukket, og API-endpointet '
    'afviser forsøg på at ændre rabatter. Valideringen sker både i præsentationslaget '
    '(UI) og i forretningslogiklaget (servicelag) for at sikre dataintegritet.'
)

doc.add_heading('3.2 Ikke-funktionelle krav', level=2)

for title_text, body_text in [
    ('Brugervenlighed',
     'Grænsefladen skal være intuitiv med et mørkt tema, så en førstegangsbruger '
     'kan gennemføre en bestilling uden vejledning. Alle handlinger giver klar feedback '
     'via visuelle indikatorer og statusmeddelelser.'),
    ('Ydeevne',
     'Menuen og ordrer skal indlæses hurtigt. '
     'Adminpanelet viser opdateret data ved navigation mellem fanerne.'),
    ('Pålidelighed',
     'Databasen opretholder dataintegriteten og sikrer korrekt lageropdatering. '
     'Systemet håndterer samtidige bestillinger korrekt.'),
    ('Sikkerhed',
     'Adgangskoder gemmes i databasen (i produktion vil man hashe med bcrypt). '
     'Kun godkendte brugere har adgang til bestillingsfunktioner, kun admin til adminfunktioner.'),
    ('Vedligeholdelse',
     'Kodebasen følger en klar trelagsarkitektur, hvilket gør det nemt at finde og '
     'ændre funktioner. Databasedesignet overholder tredje normalform.'),
]:
    add_bold_italic_label(title_text)
    add_body(body_text)

doc.add_heading('3.3 Use-case diagram', level=2)
add_body(
    'Nedenstående use-case diagram viser de to primære aktører og deres interaktioner '
    'med systemet:'
)
add_body(
    'Elev (student): En elev, der besøger kantinens websted for at se menuen og bestille mad. '
    'Dette er den primære brugertype.'
)
add_body(
    'Administrator (kantinedame): En kantinemedarbejder med udvidede rettigheder, der '
    'administrerer menu, ordrer og rabatter.'
)

add_comment_placeholder(
    'Use-case diagram (UML) der viser: '
    'Aktør "Elev" med forbindelser til: Opret konto, Log ind, Se menu (<<include>> Se allergen/halal-info), '
    'Filtrer efter kategori, Tilføj til kurv, Betal & bestil (<<include>> Generer kvitteringskode), '
    'Reserver mad til næste dag, Se ordrehistorik. '
    'Aktør "Administrator" med forbindelser til: Log ind, Administrer menu, '
    'Se alle ordrer, Marker som afhentet, Sæt rabat (kun i åbningstid), Se salgsstatistik. '
    'Aktør "System" med forbindelser til: Auto-annuller ikke-afhentede ordrer (<<trigger>> Kantinen lukker). '
    'Tegn med draw.io eller lignende værktøj.'
)

doc.add_heading('3.4 Skærmbilleder', level=2)
add_body('Herunder vises skærmbilleder af systemets vigtigste sider:')

add_comment_placeholder('Skærmbillede: Login-siden med email- og adgangskode-felter')
add_comment_placeholder('Skærmbillede: Menuvisning med menukort, kurv-bar i bunden, og statusbar (reservations- eller købsperiode)')
add_comment_placeholder('Skærmbillede: Kvitteringsmodal med den genererede kode efter et køb')
add_comment_placeholder('Skærmbillede: "Mine ordrer" med ordrehistorik og statuskoder')
add_comment_placeholder('Skærmbillede: Adminpanel - "Administrer menu" med kategorifiltre og datovalg')
add_comment_placeholder('Skærmbillede: Adminpanel - "Alle ordrer" med ordreliste og marker-som-afhentet-knapper')
add_comment_placeholder('Skærmbillede: Adminpanel - "Rabatter" med rabatjustering per vare (i åbningstid)')
add_comment_placeholder('Skærmbillede: Adminpanel - "Rabatter" uden for åbningstid (viser lukket-besked)')
add_comment_placeholder('Skærmbillede: Adminpanel - "Statistik" med linjediagram over en uges salg')
add_comment_placeholder('Skærmbillede: Adminpanel - "Statistik" med søjlediagram for enkeltdag')
add_comment_placeholder('Skærmbillede: Detaljepanel med allergen-tags og halal-badge ved hover over en vare')
add_comment_placeholder('Skærmbillede: "Mine ordrer" med en annulleret ordre (auto-annulleret ved lukketid)')

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  4. DOKUMENTATION AF PROGRAM & UDVIKLING
# ════════════════════════════════════════════════════════
doc.add_heading('4. Dokumentation af program & udvikling', level=1)

doc.add_heading('4.1 Trelagsarkitektur', level=2)
add_body(
    'BetterCantine følger en ren trelagsarkitektur, der adskiller præsentationslag, '
    'forretningslogiklag og dataadgangslag. Denne struktur sikrer overskuelighed og '
    'gør det muligt at ændre et lag uden at påvirke de andre (løs kobling). '
    'Filstrukturen afspejler direkte de tre lag:'
)

add_code_block(
    'BetterCantine/\n'
    '  public/\n'
    '    index.html          - Præsentationslag: HTML-struktur\n'
    '    styles.css           - Præsentationslag: CSS-styling\n'
    '    main.js             - Præsentationslag: klient-JavaScript + D3.js grafer\n'
    '  server.js             - Forretningslogiklag: Express routing + auto-annullering\n'
    '  cantineServicelag.js  - Forretningslogiklag: forretningsregler\n'
    '  cantineDatalag.js     - DATAADGANGSLAG: Supabase-databasekald\n'
    '  config.js             - DATAADGANGSLAG: databaseforbindelse'
)

add_bold_italic_label('Præsentationslag (klientside)')
add_body(
    'index.html definerer sidestrukturen med alle UI-sektioner (login, menu, kurv, '
    'ordrehistorik, adminpanel med fire faner). styles.css indeholder al CSS-styling med '
    'CSS custom properties (variabler) for et konsistent tema. main.js håndterer '
    'brugerinteraktioner, foretager API-kald (via fetch), opdaterer DOM\'en dynamisk '
    'og renderer statistikgrafer med D3.js (linje- og søjlediagrammer). '
    'Dette lag kender ikke til databasen - det kommunikerer kun med serveren via HTTP.'
)

add_bold_italic_label('Forretningslogiklag (server + servicelag)')
add_body(
    'server.js konfigurerer Express-webserveren, definerer API-endpoints (ruter) og kører '
    'en server-side timer (setInterval) der automatisk annullerer ikke-afhentede ordrer '
    'ved kantinens lukketid. cantineServicelag.js indeholder den centrale forretningslogik: '
    'tidsbaserede tilstande (reservationsperiode vs. købsperiode via getMinsSinceMidnight()), '
    'prisberegning (opretOrdre()), tidsvalidering af rabatter (opdaterDiscount()), '
    'automatisk ordreannullering (annullerIkkeAfhentedeOrdrer()), aggregering af '
    'salgsstatistik (hentSalgsStatistik()) og koordinering mellem præsentations- og datalag. '
    'Servicelaget importerer datalaget og kalder dets funktioner - det taler aldrig direkte med databasen.'
)

add_bold_italic_label('Dataadgangslag')
add_body(
    'config.js opretter Supabase-klientforbindelsen via miljøvariabler. '
    'cantineDatalag.js indeholder alle funktioner, der interagerer direkte med databasen: '
    'signUp, signIn, hentDagensMenu (med allergen/halal-data), opretOrdre, opdaterDiscount, '
    'annullerIkkeAfhentedeOrdrer, hentSalgsStatistik osv. '
    'Hvis vi skifter database, behøver vi kun ændre dette lag.'
)

add_comment_placeholder(
    'Diagram over trelagsarkitekturen: '
    'Øverst: Præsentationslag (index.html + styles.css + main.js + D3.js) -> pile "HTTP / Fetch" -> '
    'Forretningslogiklag (server.js + cantineServicelag.js) -> pile "Supabase JS kald" -> '
    'Dataadgangslag (cantineDatalag.js + config.js) -> pile "SQL" -> '
    'Supabase PostgreSQL Database. '
    'Vis også setInterval-timer i server.js for auto-annullering. '
    'Vis filnavne i boksene og dataflow-pile mellem lagene.'
)

add_body(
    'Denne arkitektur følger princippet om høj samhørighed (cohesion) og løs kobling '
    '(coupling): hvert lag har et klart afgrænset ansvar, og lagene kommunikerer kun '
    'via veldefinerede grænseflader. Dette gør koden nemmere at forstå, teste og vedligeholde '
    '- i tråd med KISS-princippet (Keep It Simple, Stupid).'
)

doc.add_page_break()

doc.add_heading('4.2 Database design & ER-diagram', level=2)
add_body(
    'Databasen består af seks tabeller, der tilsammen håndterer hele kantinebestillingsflowet. '
    'Designet følger tredje normalform for at undgå redundans og sikre dataintegritet.'
)

for tbl, desc in [
    ('profiles',
     'Brugerkonti med id (UUID, PK), email (UNIQUE), password, full_name, '
     'role ("student"/"admin"), created_at.'),
    ('food_categories',
     'Madkategorier med id (PK), name. En kategori har mange food items (1:M).'),
    ('food_items',
     'Madkatalog med id (PK), name, description, base_price, discount_price, '
     'category_id (FK -> food_categories), image_url, is_active, created_at.'),
    ('daily_menu',
     'Dagsmenu med id (PK), food_item_id (FK -> food_items), menu_date, '
     'total_quantity, sold_quantity, discounted_quantity. '
     'Unik constraint på (food_item_id, menu_date).'),
    ('orders',
     'Bestillinger med id (PK), user_id (FK -> profiles), status '
     '("reserved"/"paid"/"picked_up"/"cancelled"), is_reservation, receipt_code (UNIQUE), '
     'total_price, placed_at.'),
    ('order_items',
     'Ordrelinjer (krydstabel) med id (PK), order_id (FK -> orders), '
     'daily_menu_id (FK -> daily_menu), quantity, unit_price, is_discounted.'),
]:
    add_bold_italic_label(tbl)
    add_body(desc)

add_comment_placeholder(
    'ER-diagram (Entity-Relationship) der viser alle 6 tabeller med attributter og relationer: '
    'profiles --(1:M)--> orders, '
    'food_categories --(1:M)--> food_items, '
    'food_items --(1:M)--> daily_menu, '
    'daily_menu --(1:M)--> order_items, '
    'orders --(1:M)--> order_items. '
    'Vis PK og FK tydeligt. Brug dbdiagram.io, draw.io eller lignende.'
)

add_bold_italic_label('Database relationer')
add_body(
    'profiles -> orders: en-til-mange (en bruger har mange ordrer). '
    'food_categories -> food_items: en-til-mange (en kategori har mange varer). '
    'food_items -> daily_menu: en-til-mange (en vare kan være på menuen flere dage). '
    'daily_menu -> order_items: en-til-mange (et menupunkt kan bestilles i mange ordrer). '
    'orders -> order_items: en-til-mange (en ordre har mange ordrelinjer). '
    'order_items fungerer som krydstabel og implementerer mange-til-mange mellem orders og daily_menu.'
)

doc.add_heading('4.3 Analyse af tredje normalform', level=2)
add_body(
    'Databasenormalisering handler om at give databasen et godt design, der reducerer '
    'redundans og sikrer dataintegritet. For at være i 3NF skal en tabel opfylde 1NF, 2NF og 3NF.'
)

add_bold_italic_label('1. normalform (1NF): atomare værdier + primærnogle')
add_body(
    'Alle tabeller har en primærnogle og gemmer kun atomare værdier (enkeltværdier) i hvert felt. '
    'F.eks. gemmer profiles enkeltværdier for email, password og full_name - ingen arrays eller '
    'sammensatte felter. Hver række er unik via sin primærnogle.'
)

add_bold_italic_label('2. normalform (2NF): ingen delvis afhængighed')
add_body(
    'Alle tabeller bruger enkeltkolonne-primærnøgler (serielle id-felter), så der ikke '
    'kan opstå delvis afhængighed. Alle ikke-nogle-attributter afhænger af hele primærnøglen. '
    'F.eks. i order_items afhænger quantity, unit_price og is_discounted af id - ikke kun af '
    'en del af en sammensat nøgle.'
)

add_bold_italic_label('3. normalform (3NF): ingen transitive afhængigheder')
add_body(
    'Ingen ikke-nogle-attribut bestemmer en anden ikke-nogle-attribut. '
    'Eksempel: i food_items gemmes category_id som fremmednogle, men kategorinavnet står i '
    'sin egen tabel (food_categories). Hvis vi i stedet gemte kategorinavnet direkte i food_items, '
    'ville en navneændring kræve opdatering af mange rækker - det ville være en transitiv '
    'afhængighed og bryde 3NF.'
)

add_bold_italic_label('Tabel-for-tabel verifikation')
add_body(
    'profiles: email, password, full_name, role afhænger alle direkte af id. '
    'Rollen bestemmer ikke emailen. I 3NF.'
)
add_body(
    'food_items: name, description, base_price, discount_price, category_id, is_active '
    'afhænger af id. category_id er en FK, ikke en transitiv afhængighed. I 3NF.'
)
add_body(
    'daily_menu: food_item_id, menu_date, total_quantity, sold_quantity, discounted_quantity '
    'afhænger af id. Lagerbeholdning er specifik for vare+dato-kombinationen. I 3NF.'
)
add_body(
    'orders: user_id, status, receipt_code, total_price, placed_at afhænger af id. '
    'total_price gemmes som denormaliseret cache for historisk nøjagtighed - acceptabel '
    'undtagelse. Funktionelt i 3NF.'
)
add_body(
    'order_items: order_id, daily_menu_id, quantity, unit_price, is_discounted afhænger af id. '
    'Kun fremmednogler og ordrelinje-specifikke data. I 3NF.'
)

add_bold_italic_label('Fordele ved 3NF i vores system')
add_body(
    'Datakonsistens: hver oplysning findes kun ét sted. Ændring af en varepris kræver kun '
    'opdatering i food_items. Ordrehistorikken bevarer den originale pris i order_items. '
    'Opdateringsanomalier undgås, og designet er nemt at udvide med nye tabeller.'
)

doc.add_page_break()

doc.add_heading('4.4 Udvalgte use-cases med kodeeksempler', level=2)
add_body(
    'Her gennemgår vi udvalgte dele af koden i detaljer. Vi viser Supabase-kald fra '
    'datalaget og forklarer de vigtigste metoder og datastrukturer.'
)

# ── Use Case: Opret konto ──
doc.add_heading('Use Case 1: Opret ny konto', level=3)
add_body(
    'Når en bruger opretter en konto, tjekker vi først om emailen allerede findes, '
    'og opretter derefter brugeren i profiles-tabellen med rollen "student".'
)

add_body('Først tjekkes om emailen allerede eksisterer i databasen:')
add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js linje 4-8 (signUp - email check):\n'
    'const { data: existing } = await supabase\n'
    '    .from("profiles").select("id").eq("email", email).maybeSingle();\n'
    'if (existing) return { data: null, error: { message: "Email er allerede i brug" } };'
)
add_body(
    'Metoden .eq("email", email) filtrerer rækker hvor email-kolonnen matcher. '
    '.maybeSingle() returnerer een række eller null - den giver ikke fejl hvis ingen match. '
    'Vi bruger dette i stedet for .single(), som ville kaste en fejl ved tomt resultat.'
)

add_body('Derefter oprettes brugeren:')
add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js linje 10-14 (signUp - insert):\n'
    'const { data, error } = await supabase\n'
    '    .from("profiles")\n'
    '    .insert({ email, password, full_name: fullName, role: "student" })\n'
    '    .select().single();'
)
add_body(
    '.insert() tilføjer en ny række. .select().single() returnerer den nyoprettede post '
    'inklusive det autogenererede id - svarende til SQL\'s RETURNING-klausul.'
)

add_comment_placeholder(
    'Flowchart for "Opret konto":\n'
    'START -> Bruger udfylder email + password + navn -> '
    'SELECT fra profiles WHERE email = input -> '
    'Findes email? [JA -> Vis fejl "Email i brug"] / '
    '[NEJ -> INSERT i profiles -> Return brugerdata -> Vis "Konto oprettet"] -> SLUT'
)

# ── Use Case: Log ind ──
doc.add_heading('Use Case 2: Log ind', level=3)
add_body(
    'Ved login matcher vi både email og password mod profiles-tabellen.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js linje 19-29 (signIn):\n'
    'const { data, error } = await supabase\n'
    '    .from("profiles").select("*")\n'
    '    .eq("email", email).eq("password", password)\n'
    '    .maybeSingle();'
)
add_body(
    'Vi kæder to .eq()-metoder for at filtrere på både email OG password. '
    'Dette er en "authentication query". .maybeSingle() returnerer brugerobjektet '
    'eller null, som vi bruger til at bestemme om login lykkedes.'
)

add_comment_placeholder(
    'Flowchart for "Log ind":\n'
    'START -> Bruger indtaster email + password -> '
    'SELECT * FROM profiles WHERE email=? AND password=? -> '
    'Match fundet? [NEJ -> Vis fejl] / [JA -> Gem brugerdata -> '
    'Tjek rolle -> Vis app (med admin-knapper hvis admin)] -> SLUT'
)

# ── Use Case: Se menu og bestil ──
doc.add_heading('Use Case 3: Se menu og bestil', level=3)
add_body(
    'Når brugeren ser menuen, bestemmer servicelaget først om det er reservations- eller '
    'købsperiode, og henter den korrekte dags menu fra databasen.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineServicelag.js linje 23-52 (tidslogik + hentMenu):\n'
    'Vis funktionerne getMinsSinceMidnight(), erReservationsperiode() og hentMenu() '
    'der viser hvordan tidspunktet bestemmer datoen og tilstanden.'
)
add_body(
    'getMinsSinceMidnight() beregner minutter siden midnat. erReservationsperiode() '
    'returnerer true når klokken er mellem 13:45 og 07:00. hentMenu() bruger dette til at '
    'hente menuen for enten i dag (købsperiode) eller i morgen (reservation). '
    'Denne forretningslogik ligger i servicelaget - ikke i datalaget eller præsentationslaget.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js linje 32-37 (hentDagensMenu):\n'
    'Vis relationsforespørgslen med indlejrede food_items og food_categories.'
)
add_body(
    'Supabase\'s relationssyntaks, f.eks. food_items(id, name, ..., food_categories(id, name)), '
    'henter relaterede data i ét kald - svarende til SQL JOINs. '
    'Dette er effektivt og undgår N+1-forespørgselsproblemet.'
)

add_body(
    'Når brugeren bestiller, beregner servicelaget totalprisen, genererer en kvitteringskode '
    'via Supabase RPC, og opretter ordren med ordrelinjer i datalaget.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineServicelag.js linje 76-84 (opretOrdre) og '
    'cantineDatalag.js linje 87-101 (opretOrdre med insert i orders + order_items).'
)

add_comment_placeholder(
    'Flowchart for "Bestil":\n'
    'START -> Bruger klikker "Betal & bestil" -> Split kurv i rabat/normal-linjer -> '
    'Beregn totalpris -> Generer kvitteringskode (RPC) -> '
    'INSERT ordre -> INSERT ordrelinjer -> Vis kvitteringskode i modal -> SLUT'
)

# ── Use Case: Administrer menu (admin) ──
doc.add_heading('Use Case 4: Administrer menu (admin)', level=3)
add_body(
    'Når admin tilføjer en vare til menuen, tjekker vi om varen allerede er på menuen '
    'for den dato. Hvis ja, lægges antallet oveni (DRY-princip - vi genbruger eksisterende data '
    'i stedet for at skabe dubletter).'
)

add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js linje 46-69 (tilføjTilDagensMenu):\n'
    'Vis check for eksisterende item, og enten update af total_quantity eller insert.'
)
add_body(
    'Denne funktion viser et godt eksempel på CRUD-operationer: vi læser (SELECT) '
    'for at tjekke eksistens, og enten opdaterer (UPDATE) eller opretter (INSERT) baseret '
    'på resultatet. Den unikke constraint på (food_item_id, menu_date) sikrer '
    'dataintegritet på databaseniveau.'
)

# ── Use Case: Salgsstatistik (admin) ──
doc.add_heading('Use Case 5: Salgsstatistik (admin)', level=3)
add_body(
    'Statistikpanelet giver administratoren overblik over salgsdata. Flowet involverer alle tre lag: '
    'præsentationslaget renderer D3.js-grafer, servicelaget aggregerer rå data til meningsfulde '
    'metrics, og datalaget henter order_items med relationer til daily_menu, food_items og orders.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineServicelag.js linje 102-174 (hentSalgsStatistik):\n'
    'Vis hvordan rå order_items aggregeres til per-item-per-dato-objekter med metrics som '
    'total_sold, total_revenue, discounted_sold, reservations, picked_up, cancelled osv.'
)
add_body(
    'Servicelaget modtager rå data fra datalaget og bruger en Map-struktur (objekt med '
    'sammensatte nøgler: food_item_id + menu_date) til at aggregere data. For hver '
    'ordrelinje akkumuleres quantity og revenue, og der skelnes mellem rabat/fuld pris, '
    'reservation/direkte køb, og afhentet/annulleret. Denne aggregering sker server-side '
    'for at minimere datamængden til klienten.'
)

add_comment_placeholder(
    'Screenshot af kode: main.js - renderChart()-funktionen (udsnit):\n'
    'Vis D3.js-koden der bygger SVG-elementet med akser, linjer og tooltips. '
    'Forklar brugen af d3.scaleLinear(), d3.scaleBand(), d3.line() og d3.select().'
)
add_body(
    'Klientsiden bruger D3.js til at tegne grafer. For flere dage renderes et '
    'linjediagram (d3.line()) med en linje per valgt food item. For en enkelt dag renderes '
    'et søjlediagram (d3.scaleBand()). Begge diagramtyper har interaktive tooltips, '
    'der viser præcise værdier ved hover. Farvekodningen er konsistent mellem grafen, '
    'legenden og item-vælgeren.'
)

add_comment_placeholder(
    'Flowchart for "Se statistik":\n'
    'START -> Admin vælger periode/dato og metric -> '
    'Fetch /api/stats/sales?from=X&to=Y -> '
    'Servicelaget aggregerer order_items -> Return JSON -> '
    'buildItemSelector() bygger chip-liste -> renderChart() tegner D3-graf -> '
    'renderSummary() viser totaler -> SLUT'
)

# ── Use Case: Auto-annullering ──
doc.add_heading('Use Case 6: Automatisk annullering af ordrer', level=3)
add_body(
    'Systemet annullerer automatisk ikke-afhentede ordrer når kantinen lukker. '
    'Dette er et eksempel på server-side forretningslogik der kører uden brugerinteraktion.'
)

add_comment_placeholder(
    'Screenshot af kode: server.js - setInterval-timeren:\n'
    'let lastCancelDate = null;\n'
    'setInterval(async () => {\n'
    '    if (!servicelag.erKantinenAaben() && servicelag.erReservationsperiode()) {\n'
    '        const dato = servicelag.getDatoIdag();\n'
    '        if (lastCancelDate !== dato) {\n'
    '            lastCancelDate = dato;\n'
    '            await servicelag.annullerIkkeAfhentedeOrdrer();\n'
    '        }\n'
    '    }\n'
    '}, 60 * 1000);'
)
add_body(
    'Timeren kører hvert 60. sekund (setInterval med 60000 ms). Den tjekker to betingelser: '
    'at kantinen ikke er åben, og at vi er i reservationsperioden (dvs. efter kl. 13:45). '
    'lastCancelDate-variablen sikrer at annulleringen kun sker én gang per dag - '
    'uden dette ville den køre hvert minut hele aftenen.'
)

add_comment_placeholder(
    'Screenshot af kode: cantineDatalag.js - annullerIkkeAfhentedeOrdrer():\n'
    'Vis Supabase-kaldet der opdaterer status til "cancelled" for alle ordrer med '
    'status "paid" eller "reserved" på den givne dato.'
)
add_body(
    'I datalaget bruger vi Supabase\'s .in()-metode til at matche flere statusser '
    '("paid" og "reserved") i ét kald, og .gte()/.lte() til at afgrænse datoperioden. '
    'Dette er mere effektivt end at køre separate forespørgsler per status.'
)

add_comment_placeholder(
    'Flowchart for "Auto-annullering":\n'
    'START -> Timer tikker (hvert 60s) -> '
    'Er kantinen lukket? [NEJ -> Vent] / [JA -> '
    'Er det allerede annulleret i dag? [JA -> Vent] / '
    '[NEJ -> UPDATE orders SET status="cancelled" WHERE status IN ("paid","reserved") '
    'AND placed_at = i dag -> Log antal annullerede]] -> SLUT'
)

doc.add_heading('4.5 Trinvis udvikling', level=2)
add_body(
    'Projektet blev udviklet trinvist. Vi startede med at designe databaseskemaet og '
    'oprette tabellerne i Supabase. Derefter byggede vi datalaget (cantineDatalag.js) med '
    'grundlæggende CRUD-funktioner. Servicelaget (cantineServicelag.js) blev tilføjet '
    'for at adskille forretningslogikken fra databasekaldene. Til sidst byggede vi '
    'præsentationslaget (index.html + main.js) med brugergrænseflade.'
)
add_body(
    'Undervejs refaktorerede vi koden flere gange. Et eksempel er, at vi oprindeligt havde '
    'forretningslogik direkte i serverns routing-funktioner, men vi flyttede den ud i '
    'cantineServicelag.js for at overholde trelagsarkitekturen. Dette er et eksempel på '
    'refaktorering for bedre kodestruktur uden at ændre funktionaliteten.'
)
add_body(
    'Vi brugte pseudokode til at planlæg komplekse funktioner før vi skrev dem. '
    'F.eks. for bestillingsflowet:'
)
add_code_block(
    'START\n'
    '    FOR hver vare i kurven\n'
    '        IF vare har rabat OG rabatantal > 0\n'
    '            THEN opret ordrelinje med rabatpris\n'
    '            Resten: opret ordrelinje med fuld pris\n'
    '        ELSE opret ordrelinje med fuld pris\n'
    '    Beregn totalpris = sum af alle ordrelinjer\n'
    '    Generer kvitteringskode\n'
    '    Gem ordre + ordrelinjer i databasen\n'
    '    Vis kvitteringskode til brugeren\n'
    'SLUT',
    'Pseudokode for bestillingsflowet:'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  5. TEST AF PROGRAMMET
# ════════════════════════════════════════════════════════
doc.add_heading('5. Test af programmet', level=1)

add_body(
    'Vi har testet programmet på flere måder for at sikre, at det opfylder '
    'kravspecifikationen og håndterer fejlsituationer korrekt.'
)

doc.add_heading('5.1 Blackbox-test', level=2)
add_body(
    'Ved blackbox-test tester vi systemets funktionalitet udefra - vi giver input og '
    'tjekker om output matcher det forventede, uden at kigge ind i koden. '
    'Vi har testet følgende scenarier:'
)

# Blackbox test table
add_bold_italic_label('Test 1: Opret konto med gyldig data')
add_body('Input: email="test@test.dk", password="test123", navn="Test Bruger"')
add_body('Forventet resultat: Konto oprettes, besked "Konto oprettet" vises.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 2: Opret konto med eksisterende email')
add_body('Input: email="admin@bettercantine.dk" (allerede i brug)')
add_body('Forventet resultat: Fejlbesked "Email er allerede i brug".')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 3: Log ind med forkert adgangskode')
add_body('Input: email="admin@bettercantine.dk", password="forkert"')
add_body('Forventet resultat: Fejlbesked "Forkert email eller adgangskode".')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 4: Bestil varer og modtag kvitteringskode')
add_body('Input: Tilføj 2x sandwich til kurv, klik "Betal & bestil".')
add_body('Forventet resultat: Ordre oprettes, kvitteringskode (BC-XXXXXX) vises.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 5: Admin tilføjer vare til menu')
add_body('Input: Vælg dato, vælg food item, angiv antal 10, klik "Tilføj".')
add_body('Forventet resultat: Varen vises på menuen med antal 10.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 6: Rabatter uden for åbningstid')
add_body('Input: Forsøg at åbne "Rabatter"-fanen uden for kantinens åbningstid (før 07:00 eller efter 13:45).')
add_body('Forventet resultat: Besked "Rabatter kan kun sættes i kantinens åbningstid (07:00-13:45)" vises.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 7: Se salgsstatistik')
add_body('Input: Admin åbner "Statistik", vælger perioden "1 uge", metric "Antal solgt".')
add_body('Forventet resultat: Linjediagram vises med data per food item over 7 dage.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 8: Automatisk annullering af ordrer')
add_body('Input: Elev bestiller mad kl. 12:00, afhenter ikke. Kantinen lukker kl. 13:45.')
add_body('Forventet resultat: Ordren skifter automatisk status til "Annulleret" efter kl. 13:45.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_bold_italic_label('Test 9: Allergenvisning')
add_body('Input: Hold musen over en vare med registrerede allergener på menuen.')
add_body('Forventet resultat: Detaljepanel viser allergener som tags og halalstatus som badge.')
add_body('Faktisk resultat: Som forventet. Bestået.')

add_comment_placeholder(
    'Skærmbilleder af blackbox-testene - f.eks. fejlbesked ved duplikat email, '
    'kvitteringsmodal efter køb, adminmenu efter tilføjelse af vare, '
    'rabat lukket-besked, statistikgraf, annulleret ordre, allergen-detaljepanel.'
)

doc.add_heading('5.2 Idiottest', level=2)
add_body(
    'Ved idiottest prøver vi, om systemet kan håndtere meningsløst eller uventet input '
    'uden at gå ned.'
)

add_bold_italic_label('Test: Tomt login')
add_body('Input: Tryk "Log ind" uden at udfylde felter.')
add_body('Resultat: Fejlbesked "Udfyld email og adgangskode" vises. Systemet går ikke ned.')

add_bold_italic_label('Test: Bestil uden varer i kurven')
add_body('Input: Klik "Betal & bestil" med tom kurv.')
add_body('Resultat: Funktionen returnerer uden at gøre noget (guard clause: if cart.length === 0 return).')

add_bold_italic_label('Test: Tilføj flere varer end lageret')
add_body('Input: Prøv at tilføje 100 stk af en vare med 10 på lager.')
add_body('Resultat: Systemet begrænser quantity til max tilgængeligt antal (remaining).')

add_bold_italic_label('Test: Sæt rabat højere end resterende lager')
add_body('Input: Prøv at sætte rabatantal til 50 på en vare med 8 stk tilbage.')
add_body('Resultat: Alert vises med "Antal til rabat kan ikke overstige det resterende lager (8 stk)". Værdien nulstilles.')

add_bold_italic_label('Test: Statistik uden data')
add_body('Input: Vælg en periode uden salgsdata i statistikpanelet.')
add_body('Resultat: Grafen viser "Ingen salgsdata i denne periode." i stedet for en tom graf.')

doc.add_heading('5.3 Fejlhåndtering (try-catch)', level=2)
add_body(
    'Vi bruger try-catch i alle vores server-endpoints til at fange runtime-fejl '
    'og returnere en pålidelig fejlbesked til brugeren i stedet for at lade serveren gå ned.'
)

add_comment_placeholder(
    'Screenshot af kode: server.js - et eksempel på en route med try-catch, f.eks. linje 10-17:\n'
    'app.post("/api/signup", async (req, res) => {\n'
    '    try {\n'
    '        const { email, password, fullName } = req.body;\n'
    '        const result = await servicelag.signUp(email, password, fullName);\n'
    '        if (result.error) return res.status(400).json({ error: result.error.message });\n'
    '        res.json(result.data);\n'
    '    } catch (err) { console.log(err); res.status(500).json({ error: "Server fejl" }); }\n'
    '});'
)
add_body(
    'try-blokken kører den kode, der potentielt kan fejle (databasekald, netværksfejl). '
    'Hvis en fejl kastes (f.eks. database-timeout), fanges den af catch-blokken, som '
    'logger fejlen til konsollen og returnerer HTTP 500 med en generisk fejlbesked til klienten. '
    'På klient-siden fanger fetchJSON()-funktionen også netværksfejl via sin egen try-catch.'
)

add_body(
    'Der er tre typer fejl vi skal være opmærksom på: '
    'Syntaksfejl (fanges af editoren/linter før kørsel), '
    'Runtime-fejl (fanges af vores try-catch, f.eks. databasefejl), og '
    'Logiske fejl (fanges via test, f.eks. forkert prisberegning).'
)

doc.add_heading('5.4 Brugertest', level=2)
add_body(
    'Vi udførte en brugertest med en klassekammerat som testperson. Testpersonen fik '
    'følgende opgaver uden vejledning:'
)
add_body('1. Opret en ny konto')
add_body('2. Find og bestil en sandwich fra menuen')
add_body('3. Se din ordrehistorik og find kvitteringskoden')
add_body(
    'Testpersonen gennemførte alle opgaver uden problemer. Feedback: '
    'login-flowet var intuitivt, menukortene var overskuelige, og kvitteringskoden var '
    'nem at finde. Forbedringsforslag: tilføj en søgefunktion og vis allergener.'
)

add_comment_placeholder(
    'Eventuelt: skærmbillede af testpersonens skærm under brugertesten, '
    'eller noter fra observationen.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  6. KONKLUSION
# ════════════════════════════════════════════════════════
doc.add_heading('6. Konklusion', level=1)
add_body(
    'BetterCantine opfylder de opsatte kravspecifikationer. Systemet giver elever mulighed '
    'for at se menu med allergen- og halaloplysninger, bestille mad og modtage kvitteringskode. '
    'Administratorer kan styre menu, ordrer, tidsbaserede rabatter og analysere salgsdata '
    'via et interaktivt statistikpanel. Systemet annullerer automatisk ikke-afhentede ordrer '
    'ved lukketid. Databasen er designet i tredje normalform, og koden følger en '
    'ren trelagsarkitektur med klar adskillelse af ansvar.'
)
add_body(
    'Udviklingsprocessen fulgte en trinvis tilgang, hvor vi startede med databasedesign, '
    'derefter datalaget, servicelaget og til sidst præsentationslaget. Vi refaktorerede '
    'undervejs for at overholde arkitekturprincipperne - f.eks. blev tidsvalideringen af '
    'rabatter implementeret i både præsentationslaget og servicelaget for at sikre '
    'dataintegritet på flere niveauer. Koden er testet med blackbox-test, '
    'idiottest og brugertest.'
)
add_body(
    'Hvis vi havde mere tid, ville vi implementere: korrekt password-hashing med bcrypt, '
    'betalingsintegration (f.eks. MobilePay), realtidsopdateringer med Supabase Realtime, '
    'push-notifikationer til elever når deres ordre er klar, og en mobilvenlig '
    'PWA-version af applikationen. Vi ville også udvide brugertesten med flere testpersoner '
    'og mere systematisk dataindsamling.'
)

doc.add_paragraph()
add_body('Link til projektet: [Indsæt GitHub-link her]')
add_body(
    'Note: For at prøve systemet skal du oprette en konto (kræver ikke en reel email, '
    'bare noget med @ i). Der er også en admin-konto:'
)
add_body('Email: admin@bettercantine.dk')
add_body('Password: admin123')

doc.add_page_break()

# ════════════════════════════════════════════════════════
#  BILAG
# ════════════════════════════════════════════════════════
doc.add_heading('Bilag', level=1)
add_body(
    'Herunder vedlægges al kildekode med kommentarer. Koden kan også findes på GitHub '
    '(se link på forsiden).'
)

add_bold_italic_label('Bilag A: config.js (databaseforbindelse)')
add_comment_placeholder('Indsæt hele config.js med kodekommentarer')

add_bold_italic_label('Bilag B: cantineDatalag.js (dataadgangslag)')
add_comment_placeholder('Indsæt hele cantineDatalag.js med kodekommentarer')

add_bold_italic_label('Bilag C: cantineServicelag.js (forretningslogiklag)')
add_comment_placeholder('Indsæt hele cantineServicelag.js med kodekommentarer')

add_bold_italic_label('Bilag D: server.js (Express routing)')
add_comment_placeholder('Indsæt hele server.js med kodekommentarer')

add_bold_italic_label('Bilag E: public/index.html (HTML + CSS)')
add_comment_placeholder('Indsæt hele index.html med kodekommentarer')

add_bold_italic_label('Bilag F: public/main.js (klient-JavaScript)')
add_comment_placeholder('Indsæt hele main.js med kodekommentarer')

add_bold_italic_label('Bilag G: Yderligere diagrammer')
add_comment_placeholder('Eventuelt yderligere diagrammer der ikke indgår i brodteksten')

# ════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════
output_path = '/Users/5veins9/Desktop/Programming/BetterCantine/Rapport/BetterCantine_Rapport.docx'
doc.save(output_path)
print(f'Rapport gemt: {output_path}')
